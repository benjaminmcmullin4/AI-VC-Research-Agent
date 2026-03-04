"""PDF and DOCX export for PE-style market memos."""

from __future__ import annotations

import io
import re
from typing import TYPE_CHECKING

from config import COLORS, MEMO_SECTIONS

if TYPE_CHECKING:
    from schema import MarketMemo


def _memo_sections(memo: "MarketMemo") -> list[tuple[str, str]]:
    """Return ordered (heading, content) pairs for the memo."""
    return [
        ("Executive Summary", memo.executive_summary),
        ("Market Overview", memo.market_overview),
        ("TAM / SAM Sizing", memo.tam_sam_sizing),
        ("Competitive Landscape", memo.competitive_landscape),
        ("Taxonomy & Segmentation", memo.taxonomy_segmentation),
        ("Company Profiles", memo.company_profiles),
        ("Investment Thesis", memo.investment_thesis),
        ("Value Creation Angles", memo.value_creation_angles),
        ("Risks & Mitigants", memo.risks_and_mitigants),
        ("Diligence Plan", memo.diligence_plan),
        ("Sources & Methodology", memo.sources_and_methodology),
    ]


def _strip_markdown(text: str) -> str:
    """Lightly strip markdown formatting for PDF plain-text rendering."""
    text = re.sub(r"#{1,6}\s*", "", text)          # headings
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)   # bold
    text = re.sub(r"\*(.+?)\*", r"\1", text)        # italic
    text = re.sub(r"`(.+?)`", r"\1", text)          # inline code
    return text


def _sanitize_for_pdf(text: str) -> str:
    """Replace Unicode characters that fpdf2 Helvetica can't encode."""
    replacements = {
        "\u2014": "--",   # em dash
        "\u2013": "-",    # en dash
        "\u2018": "'",    # left single quote
        "\u2019": "'",    # right single quote
        "\u201c": '"',    # left double quote
        "\u201d": '"',    # right double quote
        "\u2026": "...",  # ellipsis
        "\u2022": "-",    # bullet
        "\u2192": "->",   # right arrow
        "\u2264": "<=",   # ≤
        "\u2265": ">=",   # ≥
        "\u00a0": " ",    # non-breaking space
    }
    for char, repl in replacements.items():
        text = text.replace(char, repl)
    # Fallback: replace any remaining non-latin1 chars
    return text.encode("latin-1", errors="replace").decode("latin-1")


# ── PDF Export ─────────────────────────────────────────────────────────

def generate_pdf(memo: "MarketMemo") -> bytes:
    """Generate a PDF memo and return as bytes."""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=25)
    pdf.add_page()

    # Title page
    pdf.set_font("Helvetica", "B", 24)
    pdf.set_text_color(10, 10, 10)  # primary black
    pdf.cell(0, 20, _sanitize_for_pdf(memo.title), ln=True, align="C")
    pdf.set_font("Helvetica", "", 12)
    pdf.set_text_color(119, 119, 119)  # muted
    pdf.cell(0, 10, f"Market Research Memo  |  {memo.date}", ln=True, align="C")
    pdf.ln(15)

    # Sections
    for heading, content in _memo_sections(memo):
        pdf.set_font("Helvetica", "B", 14)
        pdf.set_text_color(10, 10, 10)
        pdf.cell(0, 10, _sanitize_for_pdf(heading), ln=True)
        pdf.set_draw_color(26, 188, 156)  # teal
        pdf.line(pdf.get_x(), pdf.get_y(), pdf.get_x() + 190, pdf.get_y())
        pdf.ln(3)

        pdf.set_font("Helvetica", "", 10)
        pdf.set_text_color(44, 62, 80)  # text
        clean = _sanitize_for_pdf(_strip_markdown(content))
        pdf.multi_cell(0, 5, clean)
        pdf.ln(8)

    # Footer
    pdf.set_font("Helvetica", "I", 8)
    pdf.set_text_color(119, 119, 119)
    pdf.cell(0, 10, _sanitize_for_pdf("Internal Use Only — Confidential"), ln=True, align="C")

    return pdf.output()


# ── DOCX Export ────────────────────────────────────────────────────────

def generate_docx(memo: "MarketMemo") -> bytes:
    """Generate a DOCX memo and return as bytes."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(memo.title)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x0A, 0x0A, 0x0A)

    # Subtitle
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_para.add_run(f"Market Research Memo  |  {memo.date}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    doc.add_paragraph()  # spacer

    # Sections
    for heading, content in _memo_sections(memo):
        h = doc.add_heading(heading, level=2)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x0A, 0x0A, 0x0A)

        # Split content into paragraphs
        for para_text in content.split("\n"):
            stripped = para_text.strip()
            if not stripped:
                continue
            # Handle bullet points
            if stripped.startswith("- ") or stripped.startswith("* "):
                p = doc.add_paragraph(stripped[2:], style="List Bullet")
            elif re.match(r"^\d+\.\s", stripped):
                p = doc.add_paragraph(re.sub(r"^\d+\.\s", "", stripped), style="List Number")
            else:
                clean = _strip_markdown(stripped)
                p = doc.add_paragraph(clean)

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Internal Use Only — Confidential")
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
